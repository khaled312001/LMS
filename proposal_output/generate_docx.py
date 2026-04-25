# -*- coding: utf-8 -*-
"""
Generate plain Arabic Word document with all proposal details (no design).
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_rtl(paragraph):
    """Force the paragraph direction to RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def set_run_font(run, size=12, bold=False, name='Arial'):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_para(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_rtl(p)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(p)
    r = p.add_run(text)
    set_run_font(r, size=sizes.get(level, 12), bold=True)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(p)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def add_empty(doc):
    doc.add_paragraph()


# ====================================================================
doc = Document()

# Set default style for the whole document
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Page setup (A4)
from docx.shared import Cm
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============ TITLE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_rtl(title)
r = title.add_run('بروبوزل منصة تعليم إلكتروني متكاملة')
set_run_font(r, size=22, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_rtl(sub)
r = sub.add_run('QTD Academy — Integrated Learning Management System')
set_run_font(r, size=13, bold=False)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_rtl(meta)
r = meta.add_run('مُعدّ بواسطة: Marbrnd Agency — Marketing Agency')
set_run_font(r, size=11, bold=False)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_rtl(meta2)
r = meta2.add_run('نسخة 1.0 — أبريل 2026')
set_run_font(r, size=11)

add_empty(doc)

# ============ 1. EXECUTIVE SUMMARY ============
add_heading(doc, '١. الملخص التنفيذي', level=1)
add_para(doc, (
    'منصة تعليمية متكاملة تجمع بين التعليم المسجّل والمباشر، مبنية بالكامل على Node.js و Next.js، '
    'تخدم ثلاثة أدوار رئيسية: الطالب، المعلم، الإدارة، مع موقع تعريفي احترافي متعدد اللغات، '
    'ومنصة ويب متجاوبة بالكامل مع كل الأجهزة. تهدف المنصة إلى توفير تجربة تعليمية عصرية '
    'وآمنة وسهلة الاستخدام للطلاب والمعلمين والإدارة.'
))
add_empty(doc)

add_heading(doc, 'أبرز ملامح المشروع', level=2)
add_bullet(doc, 'ثلاثة أدوار رئيسية: طالب، معلم، إدارة، بصلاحيات منفصلة لكل دور.')
add_bullet(doc, 'منصة ويب متجاوبة بالكامل (Responsive) تعمل على كل الأجهزة.')
add_bullet(doc, 'دعم ثنائي اللغة: العربية والإنجليزية.')
add_bullet(doc, 'مدة التسليم: 3–4 أسابيع قابلة للنقصان حسب جاهزية المحتوى.')
add_bullet(doc, 'قابلية الإضافة مستقبلاً: تطبيقات موبايل أصلية عند الحاجة.')
add_empty(doc)

# ============ 2. PROJECT OBJECTIVES ============
add_heading(doc, '٢. أهداف المشروع', level=1)
add_para(doc, (
    'إطلاق منصة تعليمية جاهزة للاستخدام الفعلي خلال 3–4 أسابيع، قادرة على استقبال الطلاب '
    'وتقديم المحتوى التعليمي فور إطلاقها، مع بنية قابلة للتوسّع مستقبلاً.'
))
add_empty(doc)

add_heading(doc, 'الأهداف الرئيسية', level=2)
add_bullet(doc, 'تجربة موحّدة لكل الأدوار: دعم جميع أطراف العملية التعليمية (الطلاب، المدرسون، الإدارة) عبر نظام موحّد بصلاحيات منفصلة لكل دور.')
add_bullet(doc, 'تعليم مباشر ومسجّل: حصص مباشرة (فردية وجماعية) عبر Google Meet، إلى جانب مكتبة محتوى مسجّل بمشغّل محمي.')
add_bullet(doc, 'حماية المحتوى: حماية المحتوى من النسخ والتحميل عبر Signed URLs، Custom Video Player، وعلامة مائية ديناميكية.')
add_bullet(doc, 'ويب متجاوب مع قابلية التوسّع: منصة ويب متجاوبة بالكامل تعمل على الجوال واللاب والتابلت، مع قابلية إضافة تطبيقات موبايل مستقبلاً.')
add_bullet(doc, 'بنية قابلة للتوسّع: تجهيز البنية التحتية بحيث يمكن إضافة ميزات ومراحل جديدة دون الحاجة لإعادة بناء النظام من الصفر.')
add_bullet(doc, 'إدارة ذكية بالبيانات: لوحات تحكم مبنية على بيانات حيّة (إيرادات، تسجيلات، أداء الطلاب) لاتخاذ قرارات مدعومة بالأرقام.')
add_empty(doc)

# ============ 3. LANDING SITE ============
add_heading(doc, '٣. الموقع التعريفي — صفحة الماركتنج (Landing Site)', level=1)
add_para(doc, (
    'الواجهة الأولى للأكاديمية أمام الزوّار والعملاء، مصمّمة لتحويل الزائر إلى طالب. '
    'تدعم اللغتين العربية والإنجليزية، ومتوافقة مع جميع الأجهزة.'
))
add_empty(doc)

add_heading(doc, 'الصفحات الرئيسية', level=2)
add_bullet(doc, 'الصفحة الرئيسية: Hero Section تفاعلي، الدورات المميزة، شهادات الطلاب، إحصائيات الأكاديمية.')
add_bullet(doc, 'من نحن (About): قصة الأكاديمية والرؤية، فريق العمل والمدربين، القيم والرسالة، الإنجازات والأرقام.')
add_bullet(doc, 'قائمة الدورات: تصنيفات متعددة وفلترة، بحث متقدم، تفاصيل كل دورة، تسجيل/شراء مباشر.')
add_bullet(doc, 'المدونة (Blog): مقالات تعليمية، تصنيفات وبحث، تعليقات وتفاعل، SEO محسّن.')
add_bullet(doc, 'تواصل معنا: نموذج تواصل ذكي، تكامل WhatsApp، بريد إلكتروني وهاتف، موقع على الخريطة.')
add_bullet(doc, 'ميزات إضافية: دعم ثنائي اللغة AR/EN، FAQ، صفحات Privacy/Terms، Newsletter.')
add_empty(doc)

# ============ 4. STUDENT JOURNEY ============
add_heading(doc, '٤. رحلة الطالب — Student Journey', level=1)
add_para(doc, (
    'تجربة تعليمية متكاملة من التسجيل حتى استلام الشهادة — كل شيء في مكان واحد.'
))
add_empty(doc)

add_heading(doc, 'الرحلة الكاملة (End-to-End Journey)', level=2)
add_bullet(doc, 'التسجيل: إنشاء حساب + تفعيل البريد.')
add_bullet(doc, 'التصفح: اختيار الدورة المناسبة.')
add_bullet(doc, 'الاشتراك: الدفع عبر Stripe.')
add_bullet(doc, 'التعلم: مشاهدة الدروس المسجلة + حضور الحصص المباشرة.')
add_bullet(doc, 'الشهادة: اجتياز الاختبار + تحميل الشهادة.')
add_empty(doc)

add_heading(doc, 'المميزات المتاحة للطالب', level=2)

add_heading(doc, 'أ) التسجيل والحساب', level=3)
add_bullet(doc, 'تسجيل دخول عبر Email أو حسابات Social.')
add_bullet(doc, 'استعادة كلمة المرور عبر البريد الإلكتروني.')
add_bullet(doc, 'تفعيل الحساب برمز OTP.')
add_bullet(doc, 'ملف شخصي قابل للتعديل.')
add_bullet(doc, 'قائمة أمنيات (Wishlist) لحفظ الدورات.')

add_heading(doc, 'ب) تصفح الدورات والاشتراك', level=3)
add_bullet(doc, 'بحث متقدم بالتصنيف والسعر.')
add_bullet(doc, 'معاينة مجانية للدرس الأول.')
add_bullet(doc, 'سلة مشتريات وكوبونات خصم.')
add_bullet(doc, 'شراء دورة مفردة أو اشتراك شهري.')

add_heading(doc, 'ج) مشاهدة الدروس', level=3)
add_bullet(doc, 'مشغل فيديو محمي ضد النسخ.')
add_bullet(doc, 'سرعات تشغيل متعددة.')
add_bullet(doc, 'حفظ آخر موضع تلقائياً.')
add_bullet(doc, 'ملفات داعمة (PDF / Slides).')
add_bullet(doc, 'ترجمات (Subtitles) عند الحاجة.')

add_heading(doc, 'د) الحصص المباشرة', level=3)
add_bullet(doc, 'جدول الحصص القادمة.')
add_bullet(doc, 'الدخول بنقرة واحدة عبر Google Meet.')
add_bullet(doc, 'تذكيرات قبل بداية الحصة.')
add_bullet(doc, 'حضور حصص فردية (1-to-1) وجماعية.')

add_heading(doc, 'هـ) الاختبارات وتتبع التقدم', level=3)
add_bullet(doc, 'اختبارات بأنواع متعددة.')
add_bullet(doc, 'نتائج فورية ومراجعة الحلول.')
add_bullet(doc, 'تقدم تفصيلي لكل دورة.')
add_bullet(doc, 'شريط إنجاز (Progress Bar) تفاعلي.')

add_heading(doc, 'و) الشهادات والمجتمع', level=3)
add_bullet(doc, 'تحميل شهادة PDF تلقائياً.')
add_bullet(doc, 'QR Code للتحقق من صحة الشهادة.')
add_bullet(doc, 'مجتمع نقاش (Forum) لكل دورة.')
add_bullet(doc, 'محادثة مباشرة (Chat) مع المعلم.')
add_empty(doc)

# ============ 5. INSTRUCTOR PORTAL ============
add_heading(doc, '٥. بوابة المعلم — Instructor Portal', level=1)
add_para(doc, (
    'لوحة مستقلة تمنح المعلم رؤية شاملة على أدائه وطلّابه ودخله المالي — '
    'مثل "المكتب الشخصي" داخل الأكاديمية.'
))
add_empty(doc)

add_heading(doc, 'رحلة المعلم', level=2)
add_bullet(doc, 'التقديم (Become an Instructor).')
add_bullet(doc, 'إنشاء دورة (Modules + Lessons).')
add_bullet(doc, 'رفع محتوى (فيديو + ملفات + نصوص).')
add_bullet(doc, 'جدولة حصص عبر Google Meet.')
add_bullet(doc, 'متابعة التقارير والدخل.')
add_empty(doc)

add_heading(doc, 'المميزات المتاحة للمعلم', level=2)

add_heading(doc, 'أ) لوحة معلومات المعلم (Dashboard)', level=3)
add_bullet(doc, 'ملخص فوري لأهم المؤشرات.')
add_bullet(doc, 'عدد الطلاب المسجلين.')
add_bullet(doc, 'الدورات النشطة والمسوّدة.')
add_bullet(doc, 'الحصص القادمة وآخر النتائج.')
add_bullet(doc, 'الإيرادات والمبالغ المستحقة.')

add_heading(doc, 'ب) إدارة المحتوى والدورات', level=3)
add_bullet(doc, 'إنشاء دورة وتقسيمها لوحدات ودروس.')
add_bullet(doc, 'رفع فيديوهات / ملفات / نصوص.')
add_bullet(doc, 'ترتيب الوحدات بالسحب والإفلات (Drag & Drop).')
add_bullet(doc, 'نشر / إخفاء دورة أو درس.')
add_bullet(doc, 'إدارة المعاينة المجانية.')

add_heading(doc, 'ج) إدارة الحصص المباشرة', level=3)
add_bullet(doc, 'إنشاء اجتماعات Google Meet بنقرة واحدة.')
add_bullet(doc, 'جلسات فردية (1-to-1) وجماعية.')
add_bullet(doc, 'تعيين الطلاب لكل حصة.')
add_bullet(doc, 'تسجيل تلقائي للحصص.')
add_bullet(doc, 'تكامل مع نظام حجز الدروس (Tutor Booking).')

add_heading(doc, 'د) إدارة الاختبارات والتقييم', level=3)
add_bullet(doc, 'إنشاء اختبارات (MCQ / True-False / Essay).')
add_bullet(doc, 'مؤقت + عدد محاولات + أسئلة عشوائية.')
add_bullet(doc, 'تصحيح تلقائي ويدوي للأسئلة المقالية.')
add_bullet(doc, 'تقييم أداء الطلاب.')
add_bullet(doc, 'تقارير نتائج تفصيلية.')

add_heading(doc, 'هـ) متابعة الطلاب', level=3)
add_bullet(doc, 'قائمة كل الطلاب المسجلين في دوراتي.')
add_bullet(doc, 'نسب الإتمام ومعدلات الحضور.')
add_bullet(doc, 'تواصل مباشر (Chat / Email).')
add_bullet(doc, 'تقييم الأداء لكل طالب.')
add_bullet(doc, 'إرسال تنبيهات جماعية.')

add_heading(doc, 'و) التقارير المالية (Payouts)', level=3)
add_bullet(doc, 'نصيب المعلم من إيرادات كل دورة.')
add_bullet(doc, 'المبالغ المستحقة والمدفوعة.')
add_bullet(doc, 'سجل المعاملات المالية.')
add_bullet(doc, 'إعدادات التحويل المالي (Payout Settings).')
add_bullet(doc, 'تصدير تقارير PDF / Excel.')
add_empty(doc)

# ============ 6. ADMIN DASHBOARD ============
add_heading(doc, '٦. لوحة الإدارة — Admin Dashboard', level=1)
add_para(doc, (
    'لوحة إدارية مبنية بتقنية Next.js، سريعة وتفاعلية، تدعم اللغتين العربية والإنجليزية، '
    'وتمنح الإدارة تحكماً كاملاً في كل مكوّنات المنصة.'
))
add_empty(doc)

add_heading(doc, 'أ) إدارة المستخدمين', level=2)
add_bullet(doc, 'إضافة / تعديل / حذف الحسابات.')
add_bullet(doc, 'تعيين الصلاحيات (Roles).')
add_bullet(doc, 'الموافقة على المعلمين الجدد.')
add_bullet(doc, 'تتبع النشاط واللوجات.')
add_bullet(doc, 'حظر / تفعيل الحسابات.')

add_heading(doc, 'ب) إدارة الدورات والمحتوى', level=2)
add_bullet(doc, 'إنشاء / تعديل / حذف الدورات.')
add_bullet(doc, 'مراجعة الدورات قبل النشر.')
add_bullet(doc, 'إدارة التصنيفات (Categories).')
add_bullet(doc, 'إدارة Bootcamps.')
add_bullet(doc, 'المدونة (Blogs) والتصنيفات.')

add_heading(doc, 'ج) إدارة الحصص المباشرة', level=2)
add_bullet(doc, 'جدولة الحصص ومراجعتها.')
add_bullet(doc, 'تعيين المدرسين والطلاب.')
add_bullet(doc, 'مراجعة سجل الحضور.')
add_bullet(doc, 'إدارة Google Meet API Config.')
add_bullet(doc, 'تقارير الحصص.')

add_heading(doc, 'د) إدارة المدفوعات', level=2)
add_bullet(doc, 'متابعة كل المعاملات المالية.')
add_bullet(doc, 'بوابات الدفع (Stripe + محلي).')
add_bullet(doc, 'الاشتراكات الشهرية.')
add_bullet(doc, 'إصدار / إدارة الكوبونات.')
add_bullet(doc, 'الفواتير والمدفوعات اليدوية.')

add_heading(doc, 'هـ) التقارير والإحصاءات', level=2)
add_bullet(doc, 'عدد المستخدمين والتسجيلات.')
add_bullet(doc, 'أداء الدورات والمعلمين.')
add_bullet(doc, 'الإيرادات والمدفوعات.')
add_bullet(doc, 'Sales Report + Payout Report.')
add_bullet(doc, 'تصدير Excel / PDF.')

add_heading(doc, 'و) إعدادات النظام', level=2)
add_bullet(doc, 'Page Builder لصفحات الموقع.')
add_bullet(doc, 'إعدادات SEO و Meta Tags.')
add_bullet(doc, 'إعدادات الإشعارات.')
add_bullet(doc, 'اللغات والعملات.')
add_bullet(doc, 'API Keys (Google Meet، Stripe، ...).')
add_empty(doc)

# ============ 7. LIVE + DRM ============
add_heading(doc, '٧. الحصص المباشرة وحماية الفيديو', level=1)

add_heading(doc, 'أ) نظام البث الداخلي (Custom WebRTC)', level=2)
add_bullet(doc, 'نظام مبني داخل المنصة (WebRTC).')
add_bullet(doc, 'لا يحتاج أي برنامج خارجي.')
add_bullet(doc, 'تحكم كامل في التجربة والبيانات.')
add_bullet(doc, 'Whiteboard تفاعلي.')

add_heading(doc, 'ب) تكامل Google Meet API', level=2)
add_bullet(doc, 'إنشاء اجتماعات Google Meet تلقائياً.')
add_bullet(doc, 'حفظ الرابط داخل النظام.')
add_bullet(doc, 'جدولة عبر Google Calendar.')
add_bullet(doc, 'حصص فردية (1-to-1) وجماعية (Many-to-1).')

add_heading(doc, 'ج) نظام حماية الفيديو (DRM System)', level=2)
add_bullet(doc, 'منع التحميل: منع تحميل الفيديو عبر أي أداة أو امتداد.')
add_bullet(doc, 'تعطيل الزر الأيمن: Right-Click واختصارات الحفظ و Developer Tools.')
add_bullet(doc, 'Signed URLs: روابط مؤقتة لكل مستخدم تنتهي صلاحيتها تلقائياً.')
add_bullet(doc, 'علامة مائية (Watermark) ديناميكية باسم الطالب و IP أثناء المشاهدة.')

add_heading(doc, 'د) البنية التحتية للفيديو', level=2)
add_bullet(doc, 'تخزين على Hostinger: مساحات تخزين آمنة ومشفّرة.')
add_bullet(doc, 'Bunny.net CDN: توزيع المحتوى عالمياً مع سرعة بث عالية وجودة تكيّفية.')
add_bullet(doc, 'Custom Video Player: مشغل مخصص يدمج كل طبقات الحماية.')
add_empty(doc)

# ============ 8. EXAMS & CERTIFICATES ============
add_heading(doc, '٨. الاختبارات والشهادات', level=1)

add_heading(doc, 'أ) أنواع الأسئلة', level=2)
add_bullet(doc, 'Multiple Choice (اختيار من متعدد) مع خيارات غير محدودة وتصحيح تلقائي.')
add_bullet(doc, 'True / False (صح أو خطأ) للتقييم السريع.')
add_bullet(doc, 'Essay (مقالي) يتطلب تصحيحاً يدوياً من المعلم مع تعليقات.')

add_heading(doc, 'ب) خصائص نظام الاختبارات', level=2)
add_bullet(doc, 'Timer (مؤقت): تحديد وقت زمني والتسليم التلقائي عند انتهاء المدة.')
add_bullet(doc, 'Attempts: تحديد عدد المحاولات المسموحة لكل طالب.')
add_bullet(doc, 'Random: ترتيب عشوائي للأسئلة والخيارات لمنع الغش.')
add_bullet(doc, 'Auto + Manual: تصحيح تلقائي للموضوعي ويدوي للمقالي.')

add_heading(doc, 'ج) نظام الشهادات', level=2)
add_bullet(doc, 'توليد شهادة PDF تلقائياً فور اجتياز الطالب شروط الإتمام.')
add_bullet(doc, 'تحتوي على: اسم الطالب، اسم الدورة والمدرس، تاريخ الإتمام، المدة الكلية، الدرجة النهائية، توقيع وختم الأكاديمية، رقم تسلسلي فريد.')
add_bullet(doc, 'QR Code للتحقق: رمز فريد لكل شهادة.')
add_bullet(doc, 'صفحة تحقق عامة على الموقع.')
add_bullet(doc, 'حماية قوية ضد التزوير.')
add_bullet(doc, 'مشاركة مباشرة عبر LinkedIn.')
add_bullet(doc, 'إمكانية إلغاء الشهادة عند الحاجة.')
add_empty(doc)

# ============ 9. PAYMENTS + NOTIFICATIONS + COMMUNITY ============
add_heading(doc, '٩. المدفوعات والإشعارات والمجتمع', level=1)

add_heading(doc, 'أ) نظام المدفوعات (Stripe)', level=2)
add_bullet(doc, 'Stripe Integration: دفع آمن بالبطاقات العالمية، Apple Pay، Google Pay.')
add_bullet(doc, 'حماية ضد الاحتيال.')
add_bullet(doc, 'اشتراك شهري متجدد تلقائياً.')
add_bullet(doc, 'شراء دورة مفردة.')
add_bullet(doc, 'Team Packages (باقات الشركات والفرق).')
add_bullet(doc, 'إصدار فواتير PDF لكل عملية وإرسالها للبريد تلقائياً.')
add_bullet(doc, 'Offline Payments وكوبونات خصم.')

add_heading(doc, 'ب) نظام الإشعارات (Multi-Channel)', level=2)
add_bullet(doc, 'البريد الإلكتروني: ترحيب وتفعيل، إعادة كلمة مرور، تذكير الحصص، نتائج الاختبارات.')
add_bullet(doc, 'SMS (رسائل نصية): تذكير الحصص المباشرة، تنبيهات مهمة، OTP للتفعيل، تأكيد الدفع.')
add_bullet(doc, 'إشعارات داخل المنصة (In-App): فورية + Web Push Notifications + مركز إشعارات موحد.')

add_heading(doc, 'ج) المجتمع التعليمي (Community System)', level=2)
add_bullet(doc, 'مجموعات نقاش (Forum) لكل دورة مع مواضيع فرعية وتصويت على الإجابات.')
add_bullet(doc, 'محادثة مباشرة (Chat) بين الطالب والمعلم + محادثات جماعية + إرسال ملفات.')
add_bullet(doc, 'تقييمات ومراجعات (Reviews) لكل دورة ومعلم مع نظام Like / Dislike.')
add_empty(doc)

# ============ 10. TECH STACK ============
add_heading(doc, '١٠. التقنيات والبنية التحتية', level=1)
add_para(doc, (
    'النظام بالكامل مبني على Node.js و Next.js والتقنيات المرتبطة بهما، '
    'مما يضمن اتساق البيئة ومرونة الصيانة وسرعة التطوير المستقبلي.'
))
add_empty(doc)

add_heading(doc, 'المكونات التقنية', level=2)
add_bullet(doc, 'النظام الخلفي (Backend): Node.js + Express / NestJS — بيئة تشغيل JavaScript سريعة وقابلة للتوسع.')
add_bullet(doc, 'لوحة التحكم (Admin): Next.js + React — واجهة إدارة سريعة وتفاعلية مع SSR.')
add_bullet(doc, 'موقع الواجهة (Landing): Next.js + Tailwind — متعدد اللغات ومحسّن SEO.')
add_bullet(doc, 'بوابات الطالب والمعلم: Next.js + TypeScript — ضمن نفس الـ Stack لتقليل التعقيد.')
add_bullet(doc, 'تجربة الموبايل: Responsive Web + PWA قابلة للتثبيت، مع إمكانية بناء تطبيق أصلي مستقبلاً.')
add_bullet(doc, 'الاستضافة والدومين: Hostinger — اشتراك شهري أو سنوي يشمل الاستضافة والدومين وشهادة SSL.')
add_bullet(doc, 'الحصص المباشرة: Google Meet API — تكامل تلقائي لإنشاء الاجتماعات وإرسال الروابط وجدولتها.')
add_bullet(doc, 'المدفوعات: Stripe API — Checkout Sessions و Subscriptions و Webhooks.')
add_bullet(doc, 'قاعدة البيانات: MongoDB / PostgreSQL + Redis (كاش).')
add_bullet(doc, 'الإشعارات: Nodemailer + Twilio + FCM مع مركز تحكم موحد.')
add_bullet(doc, 'الأمان: JWT + OAuth + SSL + حماية OWASP Top 10.')
add_empty(doc)

add_heading(doc, 'تفاصيل الـ APIs', level=2)
add_heading(doc, 'Google Meet API', level=3)
add_bullet(doc, 'إنشاء اجتماعات تلقائياً.')
add_bullet(doc, 'إرسال الروابط للطلاب.')
add_bullet(doc, 'جدولة زمنية عبر Google Calendar.')
add_bullet(doc, 'تسجيل الحصة.')

add_heading(doc, 'Stripe API', level=3)
add_bullet(doc, 'Checkout Sessions للدفع الآمن.')
add_bullet(doc, 'Subscriptions (اشتراكات متجددة).')
add_bullet(doc, 'Webhooks لتلقي الأحداث.')
add_bullet(doc, 'دعم البطاقات العالمية و Apple Pay / Google Pay.')
add_bullet(doc, 'فواتير PDF تلقائية.')

add_heading(doc, 'Hostinger Hosting', level=3)
add_bullet(doc, 'اشتراك شهري أو سنوي.')
add_bullet(doc, 'يشمل الاستضافة والدومين.')
add_bullet(doc, 'شهادة SSL مجانية.')
add_bullet(doc, 'إدارة مبسطة للخوادم.')
add_empty(doc)

# ============ 11. TIMELINE ============
add_heading(doc, '١١. الجدول الزمني', level=1)
add_para(doc, 'المدة الإجمالية للمشروع: 3–4 أسابيع (قابلة للنقصان حسب جاهزية المحتوى).')
add_empty(doc)

add_bullet(doc, 'المرحلة ١ — التحليل والتخطيط السريع (الأسبوع ١): جمع المتطلبات، تصميم قاعدة البيانات، UI/UX، معمارية النظام وتوثيق الـ API.')
add_bullet(doc, 'المرحلة ٢ — Backend و Core APIs (الأسبوع ١–٢): بناء Node.js API، نظام المصادقة والصلاحيات، إدارة الدورات والمحتوى، تكامل Stripe.')
add_bullet(doc, 'المرحلة ٣ — لوحة الإدارة والموقع التعريفي (الأسبوع ٢): تطوير Admin Dashboard بـ Next.js مع Landing Site متعدد اللغات و Page Builder.')
add_bullet(doc, 'المرحلة ٤ — الحصص المباشرة وحماية الفيديو (الأسبوع ٢–٣): تكامل Google Meet API ونظام البث و Custom Video Player و Signed URLs.')
add_bullet(doc, 'المرحلة ٥ — تحسين التجربة المتجاوبة و PWA (الأسبوع ٣): ضبط التصميم المتجاوب + إعداد PWA قابلة للتثبيت.')
add_bullet(doc, 'المرحلة ٦ — الاختبار والإطلاق (الأسبوع ٣–٤): QA شامل، اختبار الأداء والأمان، تدريب الفريق، نشر الإنتاج والإطلاق الرسمي.')
add_empty(doc)

# ============ 12. CLOSING ============
add_heading(doc, '١٢. الخاتمة والخطوات التالية', level=1)
add_para(doc, (
    'كل ما ورد في هذا المقترح قابل للتنفيذ خلال 3–4 أسابيع (قابلة للنقصان حسب جاهزية المحتوى)، '
    'بفريق محترف يعمل بشفافية كاملة وتسليمات أسبوعية. نحن متحمسون لبناء منصة تكون مصدر فخر '
    'للأكاديمية وأداة نمو حقيقية لسنوات قادمة.'
))
add_empty(doc)

add_heading(doc, 'ملخص الصفقة', level=2)
add_bullet(doc, 'المدة الإجمالية: 3–4 أسابيع (قابلة للنقصان).')
add_bullet(doc, 'المنصة: ويب متجاوب + PWA (Responsive Platform).')
add_bullet(doc, 'اللغات المدعومة: العربية والإنجليزية.')
add_bullet(doc, 'قابلية التوسّع المستقبلي: إمكانية إضافة تطبيقات موبايل أصلية (iOS / Android).')
add_empty(doc)

add_heading(doc, 'الأطراف', level=2)
add_bullet(doc, 'العميل (Client): QTD Academy — أكاديمية كيو تي دي.')
add_bullet(doc, 'مقدم الخدمة (Provider): Marbrnd Agency — Marketing Agency.')

# Save
out_path = r'F:/LMS/proposal_output/QTD_Academy_LMS_Proposal.docx'
doc.save(out_path)
print('Saved:', out_path)
